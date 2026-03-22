import io
from typing import Optional, Dict, Any, List

from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.models.session import Session
from app.models.question import Question
from app.models.answer import Answer
from app.models.report_template import ReportTemplate


async def get_report_data(
    db_session: AsyncSession, session_id: int
) -> Optional[Dict[str, Any]]:
    """Собирает все данные для отчёта."""
    result = await db_session.execute(
        select(Session)
        .where(Session.id == session_id)
        .options(
            selectinload(Session.answers)
            .selectinload(Answer.question)
            .selectinload(Question.options),
            selectinload(Session.test),
        )
    )
    sess = result.scalar_one_or_none()
    if not sess:
        return None

    test = sess.test

    answers_data = []
    for answer in sess.answers:
        q = answer.question
        answer_text = ""

        if answer.selected_option_id:
            opt = next((o for o in q.options if o.id == answer.selected_option_id), None)
            answer_text = opt.text if opt else "N/A"
        elif answer.selected_option_ids:
            selected = [o.text for o in q.options if o.id in answer.selected_option_ids]
            answer_text = ", ".join(selected)
        elif answer.text_value:
            answer_text = answer.text_value
        elif answer.scale_value is not None:
            answer_text = str(answer.scale_value)

        answers_data.append({
            "question": q.text,
            "question_type": q.question_type,
            "answer": answer_text,
        })

    return {
        "test_id": test.id,
        "test_title": test.title,
        "test_description": test.description,
        "client_name": sess.client_name or "Анонимный клиент",
        "client_email": sess.client_email or "",
        "completed_at": str(sess.completed_at) if sess.completed_at else "",
        "status": sess.status,
        "results": sess.results or {},
        "answers": answers_data,
        "interpretations": {},  # заполним позже по шаблону
    }


async def _get_template(
    db_session: AsyncSession, test_id: int, report_type: str
) -> Optional[ReportTemplate]:
    result = await db_session.execute(
        select(ReportTemplate).where(
            ReportTemplate.test_id == test_id,
            ReportTemplate.report_type == report_type,
        )
    )
    return result.scalar_one_or_none()


async def _build_interpretations(
    db_session: AsyncSession,
    test_id: int,
    report_type: str,
    results: dict,
) -> dict:
    """
    Формат interpretations в шаблоне:
    {
      "metric_key": [
        {"min": 0, "max": 10, "level": "низкий", "text": "..."},
        {"min": 11, "max": 20, "level": "средний", "text": "..."}
      ]
    }
    """
    template = await _get_template(db_session, test_id, report_type)
    if not template or not template.interpretations:
        return {}

    metrics = (results or {}).get("metrics") or {}
    out: dict = {}

    for metric_key, ranges in template.interpretations.items():
        if not isinstance(ranges, list):
            continue

        score = metrics.get(metric_key)
        if score is None:
            continue

        # score может быть float/int
        try:
            score_val = float(score)
        except Exception:
            continue

        for r in ranges:
            try:
                r_min = float(r.get("min", 0))
                r_max = float(r.get("max", 999999))
            except Exception:
                continue

            if r_min <= score_val <= r_max:
                out[metric_key] = {
                    "score": score_val,
                    "level": r.get("level", ""),
                    "text": r.get("text", ""),
                    "min": r.get("min", None),
                    "max": r.get("max", None),
                }
                break

    return out


async def get_template_blocks(
    db_session: AsyncSession, test_id: int, report_type: str
) -> List[dict]:
    template = await _get_template(db_session, test_id, report_type)
    if template and template.blocks:
        # показываем только enabled
        return [b for b in template.blocks if b.get("enabled", True)]

    # fallback дефолты
    if report_type == "client":
        return [
            {"type": "total_score", "enabled": True, "title": "Общий результат", "config": {}},
            {"type": "metrics", "enabled": True, "title": "Метрики", "config": {}},
            {"type": "interpretation", "enabled": True, "title": "Интерпретация", "config": {}},
        ]
    return [
        {"type": "total_score", "enabled": True, "title": "Общий балл", "config": {}},
        {"type": "metrics", "enabled": True, "title": "Метрики", "config": {}},
        {"type": "answers_detail", "enabled": True, "title": "Детализация ответов", "config": {}},
    ]


def render_block_html(block: dict, data: dict) -> str:
    block_type = block.get("type", "")
    title = block.get("title", "")
    config = block.get("config", {}) or {}
    results = data.get("results", {}) or {}

    if block_type == "total_score":
        score = results.get("total_score", 0)
        return f"""
        <div style="background:#ffeaa7;border-radius:12px;padding:20px;text-align:center;margin:16px 0;">
          <h3 style="margin:0 0 8px;color:#2d3436;">{title}</h3>
          <div style="font-size:2.5rem;font-weight:800;color:#e74c3c;">{score}</div>
        </div>
        """

    if block_type == "metrics":
        metrics = results.get("metrics", {}) or {}
        if not metrics:
            return ""
        rows = ""
        for key, value in metrics.items():
            rows += f"""
            <div style="display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid #eee;">
              <span style="font-weight:600;color:#2d3436;">{key}</span>
              <span style="font-weight:700;color:#2980b9;">{value}</span>
            </div>
            """
        return f"""
        <div style="margin:16px 0;">
          <h3 style="color:#2d3436;margin-bottom:12px;">{title}</h3>
          {rows}
        </div>
        """

    if block_type == "computed":
        computed = results.get("computed", {}) or {}
        if not computed:
            return ""
        rows = ""
        for key, value in computed.items():
            rows += f"""
            <div style="display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid #eee;">
              <span style="font-weight:600;">{key}</span>
              <span style="font-weight:700;color:#e67e22;">{value}</span>
            </div>
            """
        return f"""
        <div style="margin:16px 0;">
          <h3 style="color:#2d3436;margin-bottom:12px;">{title}</h3>
          {rows}
        </div>
        """

    if block_type == "interpretation":
        interpretations = data.get("interpretations", {}) or {}
        if not interpretations:
            return ""
        items = ""
        for metric, info in interpretations.items():
            level = info.get("level", "")
            text = info.get("text", "")
            score = info.get("score", 0)
            items += f"""
            <div style="background:#e8f5e9;padding:16px;border-radius:10px;margin:8px 0;border-left:4px solid #4caf50;">
              <strong>{metric}</strong> (балл: {score}) {f"<em>— {level}</em>" if level else ""}<br>
              <span style="color:#2e7d32;">{text}</span>
            </div>
            """
        return f"""
        <div style="margin:16px 0;">
          <h3 style="color:#2d3436;margin-bottom:12px;">{title}</h3>
          {items}
        </div>
        """

    if block_type == "recommendations":
        text = config.get("text", "")
        if not text:
            return ""
        return f"""
        <div style="background:#e3f2fd;padding:20px;border-radius:12px;margin:16px 0;border-left:4px solid #2196f3;">
          <h3 style="color:#1565c0;margin:0 0 8px;">{title}</h3>
          <p style="color:#37474f;line-height:1.6;margin:0;">{text}</p>
        </div>
        """

    if block_type == "answers_detail":
        answers = data.get("answers", []) or []
        if not answers:
            return ""
        rows = ""
        for ans in answers:
            rows += (
                "<tr>"
                f"<td style='padding:10px;border:1px solid #ddd;'>{ans['question']}</td>"
                f"<td style='padding:10px;border:1px solid #ddd;'>{ans['answer']}</td>"
                "</tr>"
            )
        return f"""
        <div style="margin:16px 0;">
          <h3 style="color:#2d3436;margin-bottom:12px;">{title}</h3>
          <table style="width:100%;border-collapse:collapse;">
            <thead>
              <tr style="background:#3498db;color:white;">
                <th style="padding:10px;text-align:left;">Вопрос</th>
                <th style="padding:10px;text-align:left;">Ответ</th>
              </tr>
            </thead>
            <tbody>{rows}</tbody>
          </table>
        </div>
        """

    if block_type == "raw_data":
        import json
        raw = json.dumps(results, indent=2, ensure_ascii=False)
        return f"""
        <div style="margin:16px 0;">
          <h3 style="color:#2d3436;margin-bottom:12px;">{title}</h3>
          <pre style="background:#f5f5f5;padding:16px;border-radius:8px;overflow-x:auto;font-size:0.85rem;">{raw}</pre>
        </div>
        """

    return ""


async def generate_html_report(
    db_session: AsyncSession,
    session_id: int,
    report_type: str = "client",
) -> Optional[str]:
    data = await get_report_data(db_session, session_id)
    if not data:
        return None

    test_id = data["test_id"]
    results = data.get("results") or {}

    data["interpretations"] = await _build_interpretations(db_session, test_id, report_type, results)
    blocks = await get_template_blocks(db_session, test_id, report_type)

    blocks_html = ""
    for block in blocks:
        blocks_html += render_block_html(block, data)

    html = f"""
    <!DOCTYPE html>
    <html lang="ru">
    <head>
      <meta charset="UTF-8">
      <title>Отчёт: {data['test_title']}</title>
      <style>
        body {{
          font-family: 'Segoe UI', Tahoma, sans-serif;
          max-width: 800px;
          margin: 40px auto;
          padding: 0 20px;
          color: #333;
        }}
        h1 {{
          color: #2c3e50;
          border-bottom: 3px solid #3498db;
          padding-bottom: 10px;
        }}
        .info {{
          background: #f8f9fa;
          padding: 15px;
          border-radius: 8px;
          margin: 20px 0;
        }}
      </style>
    </head>
    <body>
      <h1>📋 {data['test_title']}</h1>
      <div class="info">
        <p><strong>Клиент:</strong> {data['client_name']}</p>
        {("<p><strong>Email:</strong> " + data['client_email'] + "</p>") if data['client_email'] else ""}
        <p><strong>Дата:</strong> {data['completed_at']}</p>
      </div>

      {blocks_html}

      <div style="margin-top:40px;padding-top:20px;border-top:1px solid #eee;color:#95a5a6;font-size:0.85rem;text-align:center;">
        Отчёт сгенерирован платформой ПрофДНК
      </div>
    </body>
    </html>
    """
    return html


async def generate_docx_report(
    db_session: AsyncSession,
    session_id: int,
    report_type: str = "client",
) -> Optional[io.BytesIO]:
    data = await get_report_data(db_session, session_id)
    if not data:
        return None

    test_id = data["test_id"]
    results = data.get("results") or {}

    data["interpretations"] = await _build_interpretations(db_session, test_id, report_type, results)
    blocks = await get_template_blocks(db_session, test_id, report_type)

    doc = Document()
    doc.add_heading(f'Отчёт: {data["test_title"]}', level=0)

    doc.add_heading("Информация", level=1)
    doc.add_paragraph(f'Клиент: {data["client_name"]}')
    if data["client_email"]:
        doc.add_paragraph(f'Email: {data["client_email"]}')
    doc.add_paragraph(f'Дата: {data["completed_at"]}')

    for block in blocks:
        btype = block.get("type", "")
        title = block.get("title", "")

        if btype == "total_score":
            doc.add_heading(title, level=1)
            doc.add_paragraph(f'Общий балл: {results.get("total_score", 0)}')

        elif btype == "metrics":
            metrics = results.get("metrics", {}) or {}
            if metrics:
                doc.add_heading(title, level=1)
                for k, v in metrics.items():
                    doc.add_paragraph(f"{k}: {v}")

        elif btype == "computed":
            computed = results.get("computed", {}) or {}
            if computed:
                doc.add_heading(title, level=1)
                for k, v in computed.items():
                    doc.add_paragraph(f"{k}: {v}")

        elif btype == "interpretation":
            interps = data.get("interpretations", {}) or {}
            if interps:
                doc.add_heading(title, level=1)
                for metric, info in interps.items():
                    level = info.get("level", "")
                    text = info.get("text", "")
                    score = info.get("score", 0)
                    doc.add_paragraph(f"{metric} ({score}) {('- ' + level) if level else ''}: {text}")

        elif btype == "recommendations":
            text = (block.get("config", {}) or {}).get("text", "")
            if text:
                doc.add_heading(title, level=1)
                doc.add_paragraph(text)

        elif btype == "answers_detail":
            answers = data.get("answers", []) or []
            if answers:
                doc.add_heading(title, level=1)
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Вопрос"
                hdr[1].text = "Ответ"
                for ans in answers:
                    row = table.add_row().cells
                    row[0].text = ans["question"]
                    row[1].text = ans["answer"]

        elif btype == "raw_data":
            import json
            doc.add_heading(title, level=1)
            doc.add_paragraph(json.dumps(results, indent=2, ensure_ascii=False))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer